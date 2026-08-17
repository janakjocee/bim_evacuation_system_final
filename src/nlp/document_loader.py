"""Regulation document text extraction helpers."""
from __future__ import annotations

from itertools import chain
from pathlib import Path
from zipfile import BadZipFile, ZipFile


MAX_PDF_PAGES = 2_000
MAX_REGULATION_TEXT_CHARACTERS = 2_000_000
MAX_DOCX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 100
MAX_DOCX_ARCHIVE_ENTRIES = 10_000


class RegulationDocumentError(ValueError):
    """Raised when a regulation document cannot be converted to text."""


def extract_regulation_text(path: str | Path, original_name: str | None = None) -> str:
    """Extract regulation text from TXT, MD, PDF or DOCX files."""
    file_path = Path(path)
    suffix = (Path(original_name or file_path.name).suffix or file_path.suffix).lower()

    if suffix in {".txt", ".md"}:
        return _bounded_text(
            [file_path.read_text(encoding="utf-8", errors="replace")],
            empty_message="No text could be extracted from the regulation document.",
        )
    if suffix == ".pdf":
        return _extract_pdf_text(file_path)
    if suffix == ".docx":
        return _extract_docx_text(file_path)

    raise RegulationDocumentError(
        f"Unsupported regulation file type '{suffix}'. Upload TXT, MD, PDF or DOCX."
    )


def _extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RegulationDocumentError(
            "PDF regulation upload requires the 'pypdf' package."
        ) from exc

    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            raise RegulationDocumentError("Encrypted PDF regulation documents are not supported.")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise RegulationDocumentError(
                f"The PDF exceeds the {MAX_PDF_PAGES}-page processing limit."
            )
        return _bounded_text(
            (page.extract_text() or "" for page in reader.pages),
            empty_message="No selectable text could be extracted from the PDF.",
        )
    except RegulationDocumentError:
        raise
    except Exception as exc:
        raise RegulationDocumentError("The PDF could not be read safely.") from exc


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RegulationDocumentError(
            "DOCX regulation upload requires the 'python-docx' package."
        ) from exc

    _validate_docx_archive(path)
    try:
        document = Document(str(path))
        chunks = (paragraph.text for paragraph in document.paragraphs)
        table_chunks = (
            " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            for table in document.tables
            for row in table.rows
        )
        return _bounded_text(
            chain(chunks, table_chunks),
            empty_message="No text could be extracted from the DOCX document.",
        )
    except RegulationDocumentError:
        raise
    except Exception as exc:
        raise RegulationDocumentError("The DOCX document could not be read safely.") from exc


def _bounded_text(chunks, *, empty_message: str) -> str:
    """Join extracted text while enforcing a deterministic memory boundary."""
    accepted: list[str] = []
    character_count = 0
    for chunk in chunks:
        cleaned = str(chunk).strip()
        if not cleaned:
            continue
        character_count += len(cleaned)
        if character_count > MAX_REGULATION_TEXT_CHARACTERS:
            raise RegulationDocumentError(
                "The extracted regulation text exceeds the 2,000,000-character processing limit."
            )
        accepted.append(cleaned)
    if not accepted:
        raise RegulationDocumentError(empty_message)
    return "\n\n".join(accepted)


def _validate_docx_archive(path: Path) -> None:
    """Reject malformed or excessively expanded DOCX ZIP containers."""
    try:
        with ZipFile(path) as archive:
            entries = [entry for entry in archive.infolist() if not entry.is_dir()]
    except (BadZipFile, OSError) as exc:
        raise RegulationDocumentError("The DOCX document is not a readable ZIP container.") from exc

    uncompressed = sum(entry.file_size for entry in entries)
    compressed = sum(max(entry.compress_size, 1) for entry in entries)
    if len(entries) > MAX_DOCX_ARCHIVE_ENTRIES:
        raise RegulationDocumentError("The DOCX archive contains too many entries.")
    if any(entry.flag_bits & 0x1 for entry in entries):
        raise RegulationDocumentError("Encrypted DOCX regulation documents are not supported.")
    if uncompressed > MAX_DOCX_UNCOMPRESSED_BYTES:
        raise RegulationDocumentError("The DOCX expanded content exceeds the 100 MB safety limit.")
    if uncompressed / max(compressed, 1) > MAX_DOCX_COMPRESSION_RATIO:
        raise RegulationDocumentError("The DOCX compression ratio exceeds the safety limit.")
